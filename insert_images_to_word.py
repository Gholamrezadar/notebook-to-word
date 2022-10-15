from docx import Document
from docx.shared import Inches
import os

TEMPLATE_PATH = 'a.docx'

document = Document('a.docx')

document.add_heading('Document GHD', 1)

import glob
last_part = ""
last_secc = ""
for f in glob.glob('images/**/*.png', recursive=True):
    path = f.split('\\')
    part=""
    secc=""

    if len(path) == 2:
        file_name = path[-1]
    if len(path) == 3:
        part = path[1]
        file_name = path[-1]
    if len(path) == 4:
        part = path[1]
        secc = path[2]
        file_name = path[-1]

    # print(part, secc, file_name)
    # part is not empty and secc is empty
    if part == "" and secc == "":
        document.add_picture(f)
        # document.add_paragraph(file_name[:-4])

    # part is not empty and secc is not empty
    if part != "" and secc != "":
        if last_part != part:
            last_part = part
            last_secc = ""
            document.add_page_break()
            document.add_paragraph(part)

        if last_secc != secc:
            last_secc = secc
            document.add_paragraph(secc)
            
        
        document.add_picture(f)
        # document.add_paragraph(file_name[:-4])

    # part is empty and secc is not empty
    if part != "" and secc == "":
        if last_part != part:
            last_part = part
            document.add_page_break()
            document.add_paragraph(part)

        document.add_picture(f)
        # document.add_paragraph(file_name[:-4])

    document.add_page_break()

document.save('demo.docx')

# p = document.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')

# document.add_paragraph(
#     'first item in unordered list', style='List Bullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='List Number'
# )

# document.add_picture('monty-truth.png', width=Inches(1.25))

# records = (
#     (3, '101', 'Spam'),
#     (7, '422', 'Eggs'),
#     (4, '631', 'Spam, spam, eggs, and spam')
# )

# table = document.add_table(rows=1, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'
# for qty, id, desc in records:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(qty)
#     row_cells[1].text = id
#     row_cells[2].text = desc


# document.save('demo.docx')